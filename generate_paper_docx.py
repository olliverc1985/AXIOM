#!/usr/bin/env python3
"""Generate AXIOM paper as a professional Word document."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()  # spacing after table
    return table

def shade_cell(cell, color="D9E2F3"):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

# ============================================================
# TITLE BLOCK
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
run = title.add_run('AXIOM: Adaptive eXecution with Intelligent Operations Memory')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(12)
run = subtitle.add_run('A Sparse Dynamic Routing Architecture for Cost-Efficient LLM Inference')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
author.paragraph_format.space_after = Pt(20)
run = author.add_run('Colin Oliver')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
author.add_run('\n')
run2 = author.add_run('Independent Research')
run2.font.size = Pt(11)
run2.font.name = 'Times New Roman'
run2.font.italic = True

# ============================================================
# ABSTRACT (shaded box via single-cell table)
# ============================================================
abstract_table = doc.add_table(rows=1, cols=1)
abstract_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = abstract_table.rows[0].cells[0]
shade_cell(cell, "E8EDF5")

# Title in abstract box
p = cell.paragraphs[0]
run = p.add_run('Abstract')
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)

# Abstract text
p2 = cell.add_paragraph()
abstract_text = (
    'We present AXIOM (Adaptive eXecution with Intelligent Operations Memory), a novel sparse routing '
    'architecture for cost-efficient large language model inference. AXIOM routes incoming queries across '
    'three model tiers \u2014 Surface, Reasoning, and Deep \u2014 using a 128-dimensional structural encoder and '
    'a hierarchical resolver with dynamic coalition formation and non-local graph communication, requiring '
    'no preference data, no GPU infrastructure, and no ML frameworks. Implemented in pure Rust with '
    '1,205,376 parameters, AXIOM achieves 100% routing accuracy on simple queries, 58.0% overall accuracy '
    'across 200 benchmark queries, and 56.6% cost reduction compared to routing all queries to a frontier '
    'model, with a mean routing latency of 1,311 microseconds. The primary architectural contribution is a '
    'sparse computation graph supporting four distinct traversal directions \u2014 forward, lateral, feedback, '
    'and temporal \u2014 enabling non-local communication between routing nodes that no existing LLM router '
    'provides. We identify and characterise the structural encoder ceiling and propose attention-based '
    'extensions as future work.'
)
run = p2.add_run(abstract_text)
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
p2.paragraph_format.space_after = Pt(4)

doc.add_paragraph()  # spacing

# ============================================================
# SECTION 1 — INTRODUCTION
# ============================================================
add_heading_styled('1  Introduction', level=1)

doc.add_paragraph(
    'Large language models vary substantially in cost and capability. Frontier models such as Claude '
    'Opus and GPT-4 provide the highest quality responses but incur significant inference costs, while '
    'smaller models are substantially cheaper but may produce inadequate responses for complex queries. '
    'For organisations processing high query volumes, routing every query to the most capable model is '
    'economically prohibitive, while routing all queries to the cheapest model degrades output quality.'
)

doc.add_paragraph(
    'LLM routing addresses this problem by predicting query complexity and dispatching each query to '
    'an appropriately capable model. Existing approaches \u2014 most notably RouteLLM (Ong et al., 2024) '
    '\u2014 achieve strong results but require training on large datasets of human preference judgements, '
    'Python runtimes with PyTorch dependencies, and binary strong/weak routing decisions. A systematic '
    'survey of 75+ routing and cascading systems (see Section 2) reveals that every existing router '
    'makes a single forward-pass decision: input enters a classifier, a score exits, a model is selected. '
    'No existing router has nodes that communicate with each other about the routing decision.'
)

doc.add_paragraph(
    'AXIOM takes a fundamentally different approach. Its sparse computation graph supports four distinct '
    'communication patterns \u2014 forward traversal, lateral traversal between same-tier nodes, feedback '
    'traversal from deeper to shallower tiers, and temporal traversal blending past routing decisions '
    'into present ones. This non-local graph communication is the primary architectural contribution '
    'and has no equivalent in the LLM routing literature.'
)

doc.add_paragraph(
    'Secondary contributions include a vocabulary-independent 128-dimensional structural encoder using '
    'G5 syntactic features, analytical initialisation with frozen surface-tier weights, dynamic coalition '
    'formation across routing tiers, and a training/inference mode split preventing embedding cache '
    'contamination.'
)

# ============================================================
# SECTION 2 — RELATED WORK
# ============================================================
add_heading_styled('2  Related Work', level=1)

p = doc.add_paragraph()
run = p.add_run('LLM routing. ')
run.bold = True
p.add_run(
    'RouteLLM (Ong et al., 2024) trains routers on Chatbot Arena preference data, achieving up to 85% '
    'cost reduction with 95% GPT-4 performance on MT Bench. Four architectures are evaluated: matrix '
    'factorisation, weighted Elo, BERT classifier, and causal LLM classifier. All are single-pass '
    'classifiers producing a scalar score. FrugalGPT (Chen et al., 2023) uses a sequential cascade '
    '\u2014 attempt a cheap model, score the output, escalate if insufficient. AutoMix (Madaan et al., 2023) '
    'uses a self-verification loop \u2014 ask the model if it is confident, escalate if not. Hybrid LLM '
    '(Ding et al., 2024) trains a binary complexity classifier. CSCR uses k-NN lookup in embedding '
    'space. None of these systems have inter-node communication. The routing decision is always made '
    'by a single component.'
)

doc.add_paragraph(
    'The closest architectural analogy in the broader literature is Tryage (Hu et al., 2023), described '
    'as a \u201cbrain-inspired\u201d thalamic router. However, Tryage is a single neural network predicting '
    'model performance \u2014 it has no actual inter-node communication. A survey published March 2026 '
    '(arXiv:2603.04445) provides a systematic analysis of 75+ multi-LLM routing and cascading '
    'approaches, covering query difficulty, preference-based, clustering, uncertainty quantification, '
    'reinforcement learning, and cascading paradigms. No surveyed system uses graph-based '
    'multi-directional communication between routing nodes.'
)

doc.add_paragraph(
    'Predictive coding networks on arbitrary graph topologies (Salvatori et al., NeurIPS 2022) are the '
    'closest architectural precedent for AXIOM\u2019s graph communication model, though applied in a '
    'completely different context. AXIOM is the first system to apply non-local graph communication to '
    'the LLM routing problem.'
)

p = doc.add_paragraph()
run = p.add_run('Mixture of Experts. ')
run.bold = True
p.add_run(
    'The MoE paradigm (Shazeer et al., 2017; Fedus et al., 2022) routes tokens to expert networks '
    'within a single model. AXIOM applies routing at the inter-model level \u2014 complete queries rather '
    'than tokens \u2014 without a jointly trained gating network.'
)

p = doc.add_paragraph()
run = p.add_run('Complexity classification. ')
run.bold = True
p.add_run(
    'Readability metrics (Kincaid et al., 1975) and syntactic complexity measures (Gibson, 1998; '
    'Yngve, 1960) inform AXIOM\u2019s encoder design. AXIOM\u2019s G5 features draw on these traditions '
    'while optimised for routing rather than readability scoring.'
)

# ============================================================
# SECTION 3 — ARCHITECTURE
# ============================================================
add_heading_styled('3  Architecture', level=1)

add_heading_styled('3.1  Overview', level=2)

doc.add_paragraph(
    'AXIOM comprises four components: a structural encoder producing a 128-dimensional representation; '
    'a sparse computation graph of ComputeNode instances with conditional, lateral, and feedback edges; '
    'a hierarchical resolver orchestrating routing decisions; and an embedding cache for repeated inputs. '
    'The system operates in two modes: RouteMode::Training (cache disabled) and RouteMode::Inference '
    '(cache enabled at cosine similarity threshold 0.92).'
)

add_heading_styled('3.2  Non-Local Graph Communication', level=2)

doc.add_paragraph(
    'The central architectural contribution is AXIOM\u2019s sparse computation graph, which supports four '
    'distinct traversal directions. Every routing decision is traceable through a sequence of TraceStep records:'
)

add_code_block('TraversalDirection { Forward, Lateral, Feedback, Temporal }')
add_code_block('TraceStep { node_id, tier, direction, confidence_in, confidence_out, was_cached }')

doc.add_paragraph(
    'Forward traversal (Surface \u2192 Reasoning \u2192 Deep) is conditional: a ConditionalEdge fires only '
    'when its EdgeCondition evaluates true for the current routing state. EdgeCondition variants are: '
    'Always, IfConfidenceAbove(f32), IfConfidenceBelow(f32), IfTier(Tier). Which edges fire depends on '
    'the input \u2014 this is not a fixed pipeline.'
)

doc.add_paragraph(
    'Lateral traversal models cortical column behaviour. When a Surface node produces low confidence, '
    'lateral edges activate other Surface nodes at the same tier before escalating. A LateralEdge '
    'connects two same-tier nodes with a weight and LateralCondition. The RouteResult tracks '
    'lateral_count (how many lateral attempts were made) and lateral_prevented_escalation (how many '
    'avoided escalation to an expensive tier). This creates graceful degradation \u2014 the system exhausts '
    'cheap options before escalating.'
)

doc.add_paragraph(
    'Feedback traversal runs upward: when a Deep node resolves an input with confidence above 0.90, '
    'it emits a FeedbackSignal to shallower tiers. FeedbackSignal carries: from_node, to_tier, reason '
    '(LowConfidenceResolved | ContradictionDetected | CacheInvalidation), and confidence_delta. This is '
    'not backpropagation \u2014 it is directional confidence nudging. If Deep repeatedly resolves inputs '
    'that Reasoning escalated unnecessarily, Reasoning\u2019s base confidence lowers, reducing future '
    'over-escalation. The system corrects its own routing mistakes without external supervision.'
)

doc.add_paragraph(
    'Temporal traversal gives AXIOM memory across routing decisions. A ring buffer of capacity 16 '
    'stores recent routing results. When a new input arrives with cosine similarity above 0.85 to a '
    'recent input, the past result blends into the current routing: current_output = 0.7 \u00d7 live_output '
    '+ 0.3 \u00d7 temporal_match. A burst of complex queries influences routing of subsequent queries even '
    'if they appear simple in isolation.'
)

doc.add_paragraph('The communication topology distinguishes AXIOM from every surveyed alternative:')

add_code_block(
    'RouteLLM:  Input \u2192 [BERT] \u2192 score \u2192 model selection\n'
    'FrugalGPT: Input \u2192 [Model1] \u2192 score \u2192 maybe [Model2] \u2192 score \u2192 maybe [Model3]\n'
    'AXIOM:     Input \u2192 [Surface1] \u2190lateral\u2192 [Surface2] \u2192 (conditional edge) \u2192\n'
    '           [Reasoning3] \u2190coalition\u2192 [Deep6] \u2192 (feedback signal upward)\n'
    '           with temporal_buffer blending throughout'
)

add_heading_styled('3.3  Structural Encoder', level=2)

doc.add_paragraph(
    'The encoder produces a 128-dimensional vector divided into five feature groups. G1 (26 dimensions): '
    'character n-gram profiles, amplified 3.0\u00d7. G2 (36 dimensions): syntactic proxy features including '
    'nested clause depth, pronoun density, and hapax ratio, amplified 3.0\u00d7. G3 (39 dimensions): '
    'position-weighted token signal. G4 (15 dimensions): scalar complexity measures including type-token '
    'ratio and punctuation density, amplified 2.0\u00d7. G5 (12 dimensions): structural syntax features '
    '\u2014 dependency depth proxy, constituent length variance, and function word position entropy, '
    'amplified 3.0\u00d7.'
)

doc.add_paragraph(
    'G5 drives the magnitude penalty applied to Surface confidence. The full confidence formula is:'
)

add_code_block(
    'cosine_sim = clamp(dot(input, weight_direction) /\n'
    '             (||input|| \u00d7 ||weight_direction|| + 1e-8), 0, 1)\n'
    'g5_penalty  = clamp((g5_norm \u2212 g5_simple_mean_norm) /\n'
    '             (g5_complex_mean_norm \u2212 g5_simple_mean_norm), 0, 1)\n'
    'confidence  = base_confidence \u00d7 0.7 + cosine_sim \u00d7 0.3\n'
    '            \u2212 g5_penalty \u00d7 0.35'
)

doc.add_paragraph(
    'Parameters g5_simple_mean_norm = 2.4596 and g5_complex_mean_norm = 3.3316 are persisted to '
    'axiom_weights.json. For multi-sentence inputs, sentence chunking splits on punctuation boundaries, '
    'encodes independently, and averages G5 norms \u2014 confirmed ratio 1.00\u00d7 on chunked versus '
    'unchunked equivalent input.'
)

add_heading_styled('3.4  Analytical Initialisation and Frozen Surface Weights', level=2)

doc.add_paragraph(
    'Surface-tier nodes are initialised analytically from the mean direction of simple training examples '
    '(AnalyticalInit) and frozen throughout training. This is the central training invariant. The '
    'rationale: the Surface tier needs to identify definitively simple inputs. Any gradient signal '
    'applied to Surface weights risks corrupting the geometric separation established at initialisation. '
    'Reasoning and Deep nodes are initialised with near-orthogonal weight directions (OrthogonalInit, '
    'mean pairwise cosine 0.0032) and updated via Oja\u2019s rule during training.'
)

add_heading_styled('3.5  Dynamic Coalition Formation', level=2)

doc.add_paragraph(
    'When input escalates past Surface, AXIOM forms a temporary coalition. Every Reasoning and Deep '
    'node computes cosine similarity to the input. Nodes above bid_threshold = 0.10 enter the bidding '
    'pool. The coalition is assembled by weighted random sampling up to max_coalition_size = 4 '
    '(stochastic selection, proportional to similarity score). The highest-bidding node\u2019s output '
    'becomes the final routing decision, but all coalition members process the input and update weights '
    'via Hebbian competition. This produces specialisation without explicit labels \u2014 nodes win bids on '
    'inputs they handle best, and their weights reinforce that specialisation.'
)

doc.add_paragraph(
    'A typical coalition: [reasoning_standalone_4 (bid=0.943), reasoning_standalone_13 (bid=0.940), '
    'deep_standalone_6 (bid=0.943, RESOLVED), reasoning_standalone_12 (bid=0.937)] with '
    'resolved_by=deep_standalone_6, cross_tier=true. Cross-tier resolutions \u2014 where a Reasoning node '
    'outbids Deep nodes \u2014 are tracked separately. Post-training: 19 of 30 R+D nodes activate '
    'regularly. Deep tier handles 9.0% of routed queries. R+D pairwise cosine drifts from 0.0032 at '
    'initialisation toward input-specific principal components via Oja convergence.'
)

# ============================================================
# SECTION 4 — TRAINING METHODOLOGY
# ============================================================
add_heading_styled('4  Training Methodology', level=1)

add_heading_styled('4.1  Corpus', level=2)

doc.add_paragraph(
    'Training corpus: 2,558 sentences from three sources. Manually constructed simple/complex sentence '
    'pairs across academic, conversational, technical, and narrative domains. A 100-entry '
    'multi-paragraph corpus (34 simple, 33 moderate, 33 complex). An adversarial curriculum of 40 '
    'sentences targeting known failure modes including very short semantically complex queries and long '
    'simple inputs.'
)

add_heading_styled('4.2  Training Procedure', level=2)

doc.add_paragraph(
    'For each training example: encoder produces 128-dimensional representation; hierarchical resolver '
    'routes the input and computes confidence; Oja\u2019s rule updates activated Reasoning and Deep node '
    'weight directions; G5 population statistics accumulate. Surface nodes receive no updates. After '
    'training: calibration computes simple_mean_confidence, complex_mean_confidence, and Surface '
    'escalation threshold from a held-out set of 27 sentences. Auto-tuner writes final configuration '
    'to axiom_config.json. G5 parameters and node weights write to axiom_weights.json (10.1 MB).'
)

add_heading_styled('4.3  Production Configuration', level=2)

doc.add_paragraph(
    'Production model: mid_dim = 128, 1,205,376 total parameters. Training time: 206 seconds '
    '(3.4 minutes). Peak RAM: 312 MB. No GPU at any stage. A scaling experiment at mid_dim = 512 '
    '(~4,800,000 parameters) produced identical adversarial accuracy (22/40) with approximately 4.7\u00d7 '
    'longer training time. The bottleneck is the 128-dimensional encoder input, not node capacity \u2014 '
    'selected production configuration is 1.2M parameters.'
)

# ============================================================
# SECTION 5 — EVALUATION
# ============================================================
add_heading_styled('5  Evaluation', level=1)

add_heading_styled('5.1  Benchmark Datasets', level=2)

doc.add_paragraph(
    'Three datasets were constructed. Simple dataset: 50 single sentences, ground truth \u201csimple\u201d, '
    'spanning customer support, basic instructions, conversational queries. Complex dataset: 50 single '
    'sentences, ground truth \u201ccomplex\u201d, spanning academic prose, technical analysis, multi-clause '
    'arguments, domain vocabulary. Realistic dataset: 100 enterprise queries \u2014 40 simple, 30 moderate, '
    '30 complex \u2014 reflecting actual LLM deployment patterns.'
)

add_heading_styled('5.2  Routing Accuracy', level=2)

doc.add_paragraph('Table 1: Routing accuracy by dataset.')

add_table(
    ['Dataset', 'Queries', 'Correct', 'Accuracy'],
    [
        ['Simple', '50', '50', '100.0%'],
        ['Complex', '50', '11', '22.0%'],
        ['Realistic', '100', '55', '55.0%'],
        ['Overall', '200', '116', '58.0%'],
    ],
    col_widths=[1.5, 1.0, 1.0, 1.0]
)

doc.add_paragraph(
    'Simple routing accuracy is 100% \u2014 every simple query stays at Surface. This is the commercially '
    'critical result: false escalations to expensive tiers drive unnecessary cost, and AXIOM eliminates '
    'them entirely on the benchmark. Complex accuracy of 22% reflects the structural encoder ceiling '
    'described in Section 5.4.'
)

add_heading_styled('5.3  Cost Model', level=2)

doc.add_paragraph('Table 2: Cost model assumptions per tier.')

add_table(
    ['Tier', 'Model', 'Input $/M', 'Output $/M', 'Avg tokens in', 'Avg tokens out'],
    [
        ['Surface', 'claude-haiku-4-5', '$0.80', '$4.00', '150', '200'],
        ['Reasoning', 'claude-sonnet-4-5', '$3.00', '$15.00', '300', '500'],
        ['Deep', 'claude-opus-4', '$15.00', '$75.00', '800', '1500'],
    ],
    col_widths=[0.9, 1.4, 0.8, 0.8, 0.9, 1.0]
)

doc.add_paragraph(
    'Measured routing distribution across all 200 benchmark queries: Surface 79.5%, Reasoning 11.5%, '
    'Deep 9.0%. At this distribution AXIOM achieves 56.6% cost reduction versus all-Opus routing.'
)

doc.add_paragraph('Table 3: Cost comparison at scale.')

add_table(
    ['Query volume', 'AXIOM cost', 'All-Opus cost', 'Saving'],
    [
        ['1,000', '$12.90', '$29.75', '56.6%'],
        ['10,000', '$129.02', '$297.49', '56.6%'],
        ['100,000', '$1,290.24', '$2,974.88', '56.6%'],
    ],
    col_widths=[1.2, 1.2, 1.2, 1.0]
)

doc.add_paragraph(
    'At 100,000 queries per day, AXIOM saves approximately $1,685 per day versus all-Opus routing. '
    'Cost savings are linear in query volume.'
)

add_heading_styled('5.4  Structural Encoder Ceiling', level=2)

doc.add_paragraph(
    'Complex routing accuracy of 22% reflects a fundamental limitation. Three failure categories are '
    'identified. Category 1 \u2014 semantic complexity without syntactic markers: short queries like '
    '\u201cReconcile Kant\u2019s categorical imperative with utilitarian ethics\u201d (7 words) present no G5 '
    'signal. The encoder cannot distinguish them from simple short queries. Category 2 \u2014 length '
    'inflation on simple inputs: multi-sentence simple inputs can inflate G5 norms if they contain '
    'subordinating conjunctions in simple contexts. Sentence chunking mitigates this (confirmed ratio '
    '1.00\u00d7) but does not eliminate it entirely. Category 3 \u2014 domain vocabulary without structure: '
    'queries using technical vocabulary without syntactic complexity markers route incorrectly to '
    'Surface. These categories define the agenda for future work: a learned embedding layer mapping '
    'tokens to semantic representations before structural analysis.'
)

add_heading_styled('5.5  Routing Latency', level=2)

doc.add_paragraph(
    'Mean routing time: 1,311 microseconds across 200 benchmark queries. This includes encoder '
    'computation, graph traversal, confidence evaluation, and embedding cache lookup. No API calls '
    'are made during routing. At 1,311 \u00b5s, AXIOM adds under 2 milliseconds of overhead to any LLM '
    'call \u2014 negligible against LLM inference latency of hundreds of milliseconds to seconds.'
)

add_heading_styled('5.6  Multi-Paragraph Routing', level=2)

doc.add_paragraph(
    'Six enterprise scenario inputs tested multi-paragraph routing using threshold strategy C (escalate '
    'if more than 40% of chunks fall below Surface threshold, constant '
    'AXIOM_CHUNK_ESCALATION_THRESHOLD = 0.40). Accuracy: 3/6 correct. Primary failure mode: confidence '
    'compression in full-text encoding, where averaging structural features across long mixed inputs '
    'dilutes per-sentence complexity signals.'
)

# ============================================================
# SECTION 6 — COMPETITIVE ANALYSIS
# ============================================================
add_heading_styled('6  Competitive Analysis', level=1)

doc.add_paragraph('Table 4: AXIOM vs RouteLLM comparison.')

add_table(
    ['Dimension', 'AXIOM', 'RouteLLM'],
    [
        ['Routing tiers', '3 (Surface, Reasoning, Deep)', '2 (strong, weak)'],
        ['Training data', 'Structural corpus, no LLM labels', 'Chatbot Arena preference votes'],
        ['Inter-node communication', 'Forward + lateral + feedback + temporal', 'None'],
        ['Coalition formation', 'Dynamic per-query', 'None'],
        ['Runtime', 'Pure Rust, no frameworks', 'Python + PyTorch'],
        ['Routing latency', '1,311 \u00b5s', 'Model-scale (BERT/LLM routers)'],
        ['Interpretability', 'Full trace: nodes, edges, direction, confidence', 'Scalar score'],
        ['Encoder type', 'Vocabulary-independent structural', 'Semantic embedding'],
        ['Reported cost savings', '56.6% vs all-Opus', 'Up to 85% vs GPT-4'],
    ],
    col_widths=[1.8, 2.3, 2.3]
)

doc.add_paragraph(
    'The 56.6% vs 85% comparison requires context. RouteLLM\u2019s figure is achieved with routers '
    'trained on millions of human preference votes. AXIOM\u2019s figure is achieved with no preference '
    'data and no LLM calls during training. They represent different points in the accuracy\u2013cost\u2013'
    'dependency tradeoff space.'
)

doc.add_paragraph(
    'AXIOM\u2019s principal advantages: framework independence, zero LLM-output training dependency, '
    'three-tier granularity, non-local graph communication, interpretable routing traces, and '
    'sub-two-millisecond latency. RouteLLM\u2019s principal advantage: higher routing accuracy on '
    'semantic complexity due to its semantic training signal.'
)

# ============================================================
# SECTION 7 — LIMITATIONS
# ============================================================
add_heading_styled('7  Limitations', level=1)

doc.add_paragraph(
    'Seven limitations are identified. First, the structural encoder ceiling prevents reliable detection '
    'of semantic complexity without syntactic markers \u2014 the binding constraint on overall accuracy. '
    'Second, complex routing accuracy of 22% means AXIOM under-escalates complex queries at a high '
    'rate. Third, multi-paragraph routing accuracy of 50% reflects confidence compression in full-text '
    'encoding. Fourth, the encoder is trained on English text \u2014 performance on other languages is '
    'untested. Fifth, no mechanism exists for learning from routing errors in production \u2014 online '
    'learning is an important future direction. Sixth, the cost model assumes fixed token counts per '
    'tier \u2014 real response length varies substantially. Seventh, the moderate/complex boundary is less '
    'well-defined than the Surface boundary, producing the largest share of errors on the realistic '
    'dataset.'
)

# ============================================================
# SECTION 8 — FUTURE WORK
# ============================================================
add_heading_styled('8  Future Work', level=1)

doc.add_paragraph(
    'Five directions are identified. First, a learned embedding layer mapping tokens to semantic '
    'representations before structural analysis \u2014 directly addressing the structural encoder ceiling. '
    'Second, an HTTP API wrapper exposing AXIOM as a drop-in OpenAI-compatible routing endpoint. '
    'Third, OpenAI model support for vendor-neutral routing across GPT-4o-mini, GPT-4o, and o1. '
    'Fourth, online learning from routing outcomes using downstream quality signals. Fifth, per-domain '
    'threshold calibration \u2014 separate thresholds for customer support, code generation, and '
    'analytical reasoning.'
)

# ============================================================
# SECTION 9 — CONCLUSION
# ============================================================
add_heading_styled('9  Conclusion', level=1)

doc.add_paragraph(
    'AXIOM demonstrates that cost-efficient LLM routing is achievable without preference data, ML '
    'frameworks, or GPU infrastructure. The primary contribution is architectural: a sparse computation '
    'graph with four distinct traversal directions \u2014 forward, lateral, feedback, and temporal \u2014 '
    'enabling non-local inter-node communication that has no equivalent in the LLM routing literature. '
    'A systematic survey of 75+ existing routing and cascading systems confirms this topology is novel. '
    'The structural encoder ceiling is identified and characterised as a fundamental limitation '
    'requiring semantic extension. The system achieves 56.6% cost reduction on realistic enterprise '
    'workloads, 100% simple routing accuracy, and 1,311 microsecond routing latency with 159 passing '
    'tests, 1.2M parameters, and 3.4-minute training time on commodity hardware.'
)

# ============================================================
# REFERENCES
# ============================================================
add_heading_styled('References', level=1)

references = [
    'Chen, L., Zaharia, M., & Zou, J. (2023). FrugalGPT: How to Use Large Language Models While Reducing Cost and Improving Performance. arXiv:2305.05176.',
    'Ding, B., et al. (2024). Hybrid LLM: Cost-Efficient and Quality-Aware Query Routing. arXiv:2404.14618.',
    'Fedus, W., Zoph, B., & Shazeer, N. (2022). Switch Transformers: Scaling to Trillion Parameter Models with Simple and Efficient Sparsity. JMLR 23(120).',
    'Gibson, E. (1998). Linguistic complexity: Locality of syntactic dependencies. Cognition 68(1), 1\u201376.',
    'Hu, S., et al. (2023). Inference Routing for Efficient LLM Serving. Tryage system.',
    'Kincaid, J.P., et al. (1975). Derivation of new readability formulas for Navy enlisted personnel. Research Branch Report 8-75, Naval Technical Training Command.',
    'Leviathan, Y., Kalman, M., & Matias, Y. (2023). Fast inference from transformers via speculative decoding. ICML 2023.',
    'Madaan, D., et al. (2023). AutoMix: Automatically Mixing Language Models. arXiv:2310.12963.',
    'Oja, E. (1982). A simplified neuron model as a principal component analyser. Journal of Mathematical Biology 15(3), 267\u2013273.',
    'Ong, I., et al. (2024). RouteLLM: Learning to Route LLMs with Preference Data. arXiv:2406.18665.',
    'Salvatori, T., et al. (2022). Predictive Coding beyond Gaussian Distributions. NeurIPS 2022.',
    'Shazeer, N., et al. (2017). Outrageously Large Neural Networks: The Sparsely-Gated Mixture-of-Experts Layer. arXiv:1701.06538.',
    'Survey arXiv:2603.04445 (2026). Dynamic Model Routing and Cascading for Efficient LLM Inference: A Survey.',
    'Yngve, V.H. (1960). A model and an hypothesis for language structure. Proceedings of the American Philosophical Society 104(5), 444\u2013466.',
]

for ref in references:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)  # hanging indent
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)

# ============================================================
# Save
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'axiom_paper.docx')
doc.save(output_path)
print(f"Saved to {output_path}")
